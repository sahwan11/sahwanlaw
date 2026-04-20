# Updated 2026-04-20: folded-in 5 parser improvements from verify/parse_phase2.py
# (Arabic dash variants, ordinal word-forms, L0372 schedule numbering,
#  L5418 EN word-form mapping, HTML whitespace collapse)
"""
parse_articles.py
-----------------
Parses every law file under C:\\Users\\sahwa\\Documents\\projects\\bahrain-laws\\
{Arabic,English} into per-article rows, keyed by LLOC slug.

Outputs scripts/_parsed_articles.json:

  {
    "K0915": {
      "notes": [...],
      "articles": [
        {
          "article_number": "1",
          "sort_order": 1000,
          "chapter": "Chapter One: ...",
          "section": null,
          "heading_ar": "...",
          "heading_en": "...",
          "text_ar": "...",
          "text_en": "..."
        },
        ...
      ]
    },
    ...
  }

Per-slug articles are MATCHED across AR + EN by article_number. If an article
exists in one language but not the other, the missing side is left null.

Regex fallback: if a file can't be split into articles at all, one row with
article_number = "UNPARSED" is produced carrying the full body.

Run:
    python scripts/parse_articles.py
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document

LAWS_ROOT = Path(r"C:\Users\sahwa\Documents\projects\bahrain-laws")
OUT_FILE = Path(__file__).resolve().parent / "_parsed_articles.json"

SLUG_RE = re.compile(r"^([A-Z]+\d+)_")

# =====================================================================
# Article-delimiter regex
# =====================================================================
# Arabic — match "مادة", "المادة", optionally preceded by "رقم", then a
# parenthesised / dash-wrapped / bare number, then an optional "مكرر/مكرراً/مكررا".
# Normalise whitespace and Arabic-Indic digits before matching.
ARABIC_DIGIT_MAP = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")

# PORTED FROM parse_phase2.py (improvement #1: Arabic dash variants)
# Covers ASCII hyphen-minus plus the Unicode dash block U+2010..U+2015
# (figure dash, en dash, em dash, horizontal bar). Some LLOC .docx files use
# forms like "مادة –1-" or "مادة - 1 -" which the old bare/parenthesised-only
# regex missed.
DASH_CHARS = r"[\-\u2010\u2011\u2012\u2013\u2014\u2015\u2212]"

RE_AR_ARTICLE = re.compile(
    r"(?:^|\n)\s*"
    r"(?:ال)?م[اـ]*\s*د[اـ]*ة"          # مادة or المادة — tolerate tatweel (ـ) and whitespace between letters
    r"(?:\s+رقم)?"                       # optional "رقم"
    r"\s*"
    r"(?:"
        # (N) or (N مكرراً) or (N مكرراً 1) or (N فقرة ك) —
        # suffix text can live INSIDE the parentheses. We capture the suf
        # via `suf_in` so later handling is uniform with the trailing form.
        # FIX 2026-04-20: added `suf_in` capture to catch L2101 amendment-style
        # "مادة (168 مكرراً)", "مادة (361 فقرة ك)", "مادة (241 مكرراً 1)" which
        # the old regex missed (swallowed Art 120 for ~9KB).
        r"\(\s*(?P<n1>\d+)(?P<suf_in>\s*(?:مكرراً|مكررا|مكرر)?"
        r"(?:\s+\d+)?"                   # optional inner ordinal like "مكرراً 1"
        r"(?:\s+فقرة\s*[^)]*)?"         # optional "فقرة X / فقرة (أ)..."
        r")?\s*\)"
        r"|"
        rf"{DASH_CHARS}\s*(?P<n2>\d+)\s*{DASH_CHARS}"  # -N- / –N– / - N -
        r"|"
        r"(?P<n3>\d+)"                   # bare N
    r")"
    r"\s*(?P<suf>مكرراً|مكررا|مكرر)?"    # bis/ter markers (optional, OUTSIDE parens)
    r"\s*:?\s*",
    re.MULTILINE,
)

# PORTED FROM parse_phase2.py (improvement #2: Arabic ordinal word-forms)
# A few laws (notably K3210 and some issuing-preambles) number articles with
# ordinal words rather than digits: "المادة الأولى", "المادة الثانية", ...
# Extended beyond phase2's cap of 20 to include compound ordinals through 40.
AR_ORDINALS = {
    "الأولى": 1, "الاولى": 1,
    "الثانية": 2,
    "الثالثة": 3,
    "الرابعة": 4,
    "الخامسة": 5,
    "السادسة": 6,
    "السابعة": 7,
    "الثامنة": 8,
    "التاسعة": 9,
    "العاشرة": 10,
    "الحادية عشرة": 11, "الحاديةعشرة": 11,
    "الثانية عشرة": 12,
    "الثالثة عشرة": 13,
    "الرابعة عشرة": 14,
    "الخامسة عشرة": 15,
    "السادسة عشرة": 16,
    "السابعة عشرة": 17,
    "الثامنة عشرة": 18,
    "التاسعة عشرة": 19,
    "العشرون": 20,
    # Extended 21-40 (compound form: "<unit> و<tens>")
    "الحادية والعشرون": 21,
    "الثانية والعشرون": 22,
    "الثالثة والعشرون": 23,
    "الرابعة والعشرون": 24,
    "الخامسة والعشرون": 25,
    "السادسة والعشرون": 26,
    "السابعة والعشرون": 27,
    "الثامنة والعشرون": 28,
    "التاسعة والعشرون": 29,
    "الثلاثون": 30,
    "الحادية والثلاثون": 31,
    "الثانية والثلاثون": 32,
    "الثالثة والثلاثون": 33,
    "الرابعة والثلاثون": 34,
    "الخامسة والثلاثون": 35,
    "السادسة والثلاثون": 36,
    "السابعة والثلاثون": 37,
    "الثامنة والثلاثون": 38,
    "التاسعة والثلاثون": 39,
    "الأربعون": 40, "الاربعون": 40,
}
_ar_ord_alt = "|".join(sorted(AR_ORDINALS.keys(), key=len, reverse=True))
RE_AR_ARTICLE_WORD = re.compile(
    rf"(?:^|\n)\s*المادة\s+({_ar_ord_alt})\b",
    re.MULTILINE,
)

# English forms observed:  "Article (1)"  "Article 1"  "Article -1-"  "Article 1:"
RE_EN_ARTICLE = re.compile(
    r"(?:^|\n)\s*Article\s*[\(\-]?\s*(\d+)\s*[\)\-]?\s*(bis|ter|quater)?\s*[:.]?\s*",
    re.IGNORECASE | re.MULTILINE,
)

# PORTED FROM parse_phase2.py (improvement #4: L5418 EN word-form mapping)
# L5418 and some translated laws spell article numbers as words:
# "Article One", "Article Twenty-One", "Article Forty". Map word-forms to
# digits and merge with the numeric stream.
WORD2NUM = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7,
    "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12, "thirteen": 13,
    "fourteen": 14, "fifteen": 15, "sixteen": 16, "seventeen": 17, "eighteen": 18,
    "nineteen": 19, "twenty": 20, "thirty": 30, "forty": 40,
}
_compounds = {}
for _t, _tv in [("twenty", 20), ("thirty", 30)]:
    for _u, _uv in [("one", 1), ("two", 2), ("three", 3), ("four", 4), ("five", 5),
                    ("six", 6), ("seven", 7), ("eight", 8), ("nine", 9)]:
        _compounds[f"{_t}-{_u}"] = _tv + _uv
        _compounds[f"{_t} {_u}"] = _tv + _uv
WORD2NUM.update(_compounds)
_word_alt = "|".join(sorted(WORD2NUM.keys(), key=len, reverse=True))
RE_EN_ARTICLE_WORD = re.compile(
    rf"(?:^|\n)\s*Article\s*[\(\-]?\s*({_word_alt})\s*[\)\-]?\s*[:.]?\s*",
    re.IGNORECASE | re.MULTILINE,
)

# Chapter/section markers for lightweight hierarchy reconstruction
RE_AR_CHAPTER = re.compile(r"^\s*(?:الباب|الفصل)\s+[^\n]*", re.MULTILINE)
RE_EN_CHAPTER = re.compile(r"^\s*(?:Chapter|Section|Part|Title)\s+[^\n]*", re.IGNORECASE | re.MULTILINE)


def slug_from_filename(name: str) -> Optional[str]:
    m = SLUG_RE.match(name)
    return m.group(1) if m else None


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


LLOC_DROP_SELECTORS = [
    "header", "footer", "nav", "script", "style", "noscript",
    ".header", ".footer", ".navbar", ".menu", ".sidebar", ".breadcrumb",
    ".search", ".social", ".language", ".top-bar",
    "#header", "#footer", "#nav", "#menu", "#sidebar",
]


def extract_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")

    # Strip scripts/styles outright (in addition to selector-based drop below)
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    main = (
        soup.find("div", id="legislationtextcontent")
        or soup.find("div", class_="legislationtextcontent")
        or soup.find("div", id="legislationtext")
        or soup.find("div", class_="legislationtext")
        or soup.find("main")
        or soup.find("article")
        or soup.body
        or soup
    )
    for sel in LLOC_DROP_SELECTORS:
        for el in main.select(sel):
            el.decompose()

    # PORTED FROM parse_phase2.py (improvement #5: HTML whitespace collapse)
    # Walk block-level tags and collapse their text. LLOC HTML often renders
    # each Arabic word on its own line (one-word-per-<br> / nested <span>),
    # which broke the "(?:^|\n)" anchor on RE_AR_ARTICLE. Joining within each
    # block with spaces — then joining blocks with newlines — restores natural
    # paragraph structure that the article regex expects.
    blocks = main.find_all(["p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6"])
    if blocks:
        texts = []
        for b in blocks:
            # Skip containers that have block-level descendants (avoid double-counting)
            if b.find(["p", "div", "li"], recursive=False):
                continue
            t = b.get_text(" ", strip=True)
            t = re.sub(r"\s+", " ", t).strip()
            if t:
                texts.append(t)
        text = "\n".join(texts)
    else:
        text = main.get_text("\n", strip=True)

    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def normalize_arabic_digits(s: str) -> str:
    return s.translate(ARABIC_DIGIT_MAP)


def _suffix_offset(suffix: Optional[str]) -> int:
    if not suffix:
        return 0
    s = suffix.strip().lower()
    if s in {"bis", "مكرر", "مكررا", "مكرراً"}:
        return 100
    if s in {"ter"}:
        return 200
    if s in {"quater"}:
        return 300
    # FIX 2026-04-20: inner-paren compound suffixes ("مكرراً 1", "فقرة ك")
    # land here. Give them a stable-but-distinct offset so each variant
    # of the same base article sorts after the plain base and after
    # the simple "مكرراً" variant.
    if "مكرر" in s and re.search(r"\d", s):
        # e.g. "مكرراً 1" — bump by the trailing integer so 1→101, 2→102
        m_n = re.search(r"\d+", s)
        return 100 + int(m_n.group(0)) if m_n else 150
    if "فقرة" in s:
        return 400
    return 500


def split_articles_arabic(text: str) -> list[dict]:
    """Split Arabic text on article markers. Returns list of dicts:
    {article_number, sort_order, suffix_raw, text_ar}

    Combines numeric article-markers (RE_AR_ARTICLE) with ordinal word-forms
    (RE_AR_ARTICLE_WORD). If a law uses BOTH, word-form matches are treated as
    preamble-style articles (sort before numerics, labelled "Preamble-N") so
    they don't collide with the digit stream.
    """
    text = normalize_arabic_digits(text)

    # Collect numeric matches (unified named-group pattern)
    num_events: list[tuple[int, int, int, Optional[str]]] = []
    for m in RE_AR_ARTICLE.finditer(text):
        base = int(m.group("n1") or m.group("n2") or m.group("n3"))
        # Suffix may be outside OR inside the parens (see FIX 2026-04-20).
        suffix_out = (m.group("suf") or "").strip() or None
        suffix_in_raw = (m.group("suf_in") or "").strip() if m.groupdict().get("suf_in") else ""
        suffix_in: Optional[str] = None
        if suffix_in_raw:
            # Normalise: collapse whitespace; treat "مكرراً", "مكررا", "مكرر" as bis.
            si = re.sub(r"\s+", " ", suffix_in_raw).strip()
            if si:
                suffix_in = si
        suffix = suffix_out or suffix_in
        num_events.append((m.start(), m.end(), base, suffix))

    # Collect word-ordinal matches, de-duplicated against numeric hits that
    # start within a few chars (some files write "المادة الأولى (1)").
    word_events: list[tuple[int, int, int, Optional[str]]] = []
    for m in RE_AR_ARTICLE_WORD.finditer(text):
        word = m.group(1)
        base = AR_ORDINALS.get(word, 0)
        if base == 0:
            continue
        if any(abs(m.start() - s) < 5 for s, _, _, _ in num_events):
            continue
        word_events.append((m.start(), m.end(), base, None))

    has_numeric = len(num_events) > 0

    # Combine into a single event stream tagged with how to name each article.
    events: list[tuple[int, int, str, int, Optional[str]]] = []  # (start, end, article_number, sort_order, suffix_raw)
    for st, en, base, suf in num_events:
        an = str(base) + (f" {suf}" if suf else "")
        so = base * 1000 + _suffix_offset(suf)
        events.append((st, en, an, so, suf))
    for st, en, base, suf in word_events:
        if has_numeric:
            # Side-channel: preamble-style articles before the main numeric sequence
            an = f"Preamble-{base}"
            so = -1000 + base
        else:
            an = str(base)
            so = base * 1000
        events.append((st, en, an, so, suf))

    events.sort(key=lambda x: x[0])
    if not events:
        return []

    # FIX 2026-04-20: TOC-tail truncator. LLOC consolidated docx sometimes
    # append a Table of Contents / فهرس after the last article (observed on
    # L1901 Civil Code Art 1054 which absorbed ~11KB of TOC rows). We detect
    # the last article's body only and truncate at the first TOC marker,
    # allowing Arabic tatweel stretch in the word "الفهرس".
    TOC_TAIL_RE = re.compile(r"\n\s*(?:ال)?ف[هـ]*ر[سـ]*\s*\n", re.MULTILINE)

    results: list[dict] = []
    for i, (mstart, mend, an, so, suf) in enumerate(events):
        start = mend
        end = events[i + 1][0] if i + 1 < len(events) else len(text)
        body = text[start:end].strip()
        # Only consider TOC truncation on the LAST article (no subsequent
        # article delimiter to naturally end it).
        if i == len(events) - 1:
            m_toc = TOC_TAIL_RE.search(body)
            if m_toc and m_toc.start() > 20:  # keep non-trivial prefix
                body = body[:m_toc.start()].rstrip()
        results.append({
            "article_number": an,
            "sort_order": so,
            "suffix_raw": suf,
            "text_ar": body,
        })
    return results


def split_articles_english(text: str) -> list[dict]:
    """Split English text on article markers. Merges numeric ("Article 1")
    and word-form ("Article One", "Article Twenty-One") matches; word-forms
    are skipped if a numeric match lands within 5 chars of the same offset."""
    num_matches = list(RE_EN_ARTICLE.finditer(text))
    word_matches = list(RE_EN_ARTICLE_WORD.finditer(text))

    merged: list[tuple[str, re.Match]] = []
    covered_starts: set[int] = set()
    for m in num_matches:
        merged.append(("num", m))
        covered_starts.add(m.start())
    for m in word_matches:
        if not any(abs(m.start() - s) < 5 for s in covered_starts):
            merged.append(("word", m))
    merged.sort(key=lambda x: x[1].start())

    if not merged:
        return []

    # FIX 2026-04-20: EN appendix-tail truncator. K0915 EN (UNCITRAL Model
    # Law docx) has a "Part Three / Recommendation regarding ..." annex
    # appended after Art 36 which had no subsequent Article marker and
    # therefore got absorbed (~50KB blob). Generic tail-truncator for last
    # article only; only triggers on fairly strong boundary words.
    EN_TAIL_RE = re.compile(
        r"\n\s*(?:Part\s+(?:Two|Three|IV|V)|"
        r"EXPLANATORY\s+NOTE|"
        r"Recommendation\s+regarding|"
        r"Resolutions?\s+adopted\s+by|"
        r"Annex\s+[IVX]+\b)",
        re.IGNORECASE,
    )

    results: list[dict] = []
    for i, (kind, m) in enumerate(merged):
        if kind == "num":
            base = int(m.group(1))
            suffix = (m.group(2) or "").strip() or None
        else:
            word = m.group(1).lower().replace(" ", "-")
            base = WORD2NUM.get(word, WORD2NUM.get(word.replace("-", " "), 0))
            if base == 0:
                continue
            suffix = None
        start = m.end()
        end = merged[i + 1][1].start() if i + 1 < len(merged) else len(text)
        body = text[start:end].strip()
        # Trim annex tail on the LAST article only
        if i == len(merged) - 1:
            m_tail = EN_TAIL_RE.search(body)
            if m_tail and m_tail.start() > 20:
                body = body[:m_tail.start()].rstrip()
        article_number = str(base) + (f" {suffix}" if suffix else "")
        sort_order = base * 1000 + _suffix_offset(suffix)
        results.append({
            "article_number": article_number,
            "sort_order": sort_order,
            "suffix_raw": suffix,
            "text_en": body,
        })
    return results


def track_chapters_ar(full_text: str) -> list[tuple[int, str]]:
    """Returns a list of (offset_in_text, chapter_heading) sorted by offset."""
    return [(m.start(), m.group(0).strip()) for m in RE_AR_CHAPTER.finditer(full_text)]


def track_chapters_en(full_text: str) -> list[tuple[int, str]]:
    return [(m.start(), m.group(0).strip()) for m in RE_EN_CHAPTER.finditer(full_text)]


def chapter_at(offset: int, chapters: list[tuple[int, str]]) -> Optional[str]:
    """Return the most-recent chapter heading before the given offset."""
    current = None
    for off, heading in chapters:
        if off <= offset:
            current = heading
        else:
            break
    return current


def _extract_l0372_schedule_items(body: str) -> list[dict]:
    """PORTED FROM parse_phase2.py (improvement #3: L0372 schedule numbering).

    Court Fees Law (L0372) has two fee schedules (جدول رقم (1), جدول رقم (2))
    whose items live OUTSIDE the main `مادة` numbering. Each item is written
    as "N- <heading>\\n<body>". We emit them as pseudo-articles with
    article_number = "Schedule-T-N" (T = table index, N = item number) and
    sort_order in the 1.0e6+ range so they land after the main articles in
    stable-sort order.
    """
    body_n = normalize_arabic_digits(body)
    table_starts = [m.start() for m in re.finditer(r"جدول\s*رقم", body_n)]
    seen: set[str] = set()
    rows: list[dict] = []
    for ti, tstart in enumerate(table_starts):
        tend = table_starts[ti + 1] if ti + 1 < len(table_starts) else len(body_n)
        sched_block = body_n[tstart:tend]
        sched_items = list(re.finditer(
            r"(?:^|\n)\s*(\d+)\s*" + DASH_CHARS + r"\s+([^\n]{3,})",
            sched_block, re.MULTILINE,
        ))
        for i, m in enumerate(sched_items):
            num = int(m.group(1))
            if num > 50:  # defensive: real schedule items don't exceed this
                continue
            heading = m.group(2).strip()
            start = m.end()
            end = sched_items[i + 1].start() if i + 1 < len(sched_items) else len(sched_block)
            item_body = sched_block[start:end].strip()
            key = f"T{ti+1}-{num}"
            if key in seen:
                continue
            seen.add(key)
            rows.append({
                "article_number": f"Schedule-{ti+1}-{num}",
                "sort_order": 1_000_000 + (ti + 1) * 10_000 + num * 100,
                "suffix_raw": None,
                "text_ar": item_body,
                "_chapter_ar": f"جدول رقم ({ti+1})",
                "_heading_ar": heading,
            })
    return rows


def parse_slug(slug: str, ar_text: Optional[str], en_text: Optional[str]) -> tuple[list[dict], list[str]]:
    """Returns (articles, notes). Articles merged by article_number across AR/EN."""
    notes: list[str] = []

    ar_articles = split_articles_arabic(ar_text) if ar_text else []
    en_articles = split_articles_english(en_text) if en_text else []

    # Recover chapter/section hierarchy per language.
    ar_chapters = track_chapters_ar(normalize_arabic_digits(ar_text or ""))
    en_chapters = track_chapters_en(en_text or "")

    # Attach chapter headings by matching article start offset in the original full text
    # (coarse approximation — good enough for MVP).
    if ar_text:
        ar_norm = normalize_arabic_digits(ar_text)
        cursor = 0
        for art in ar_articles:
            idx = ar_norm.find(art["text_ar"][:80], cursor) if art["text_ar"] else -1
            if idx < 0:
                idx = cursor
            art["_ar_offset"] = idx
            cursor = idx
            art["chapter_ar"] = chapter_at(idx, ar_chapters)

    if en_text:
        cursor = 0
        for art in en_articles:
            idx = en_text.find(art["text_en"][:80], cursor) if art["text_en"] else -1
            if idx < 0:
                idx = cursor
            art["_en_offset"] = idx
            cursor = idx
            art["chapter_en"] = chapter_at(idx, en_chapters)

    # L0372 schedule fallback: append fee-schedule items as Schedule-T-N rows
    # AFTER main articles. Done here (not inside split_articles_arabic) because
    # the slug is the trigger and the output schema (heading_ar) is set at
    # merge time below.
    schedule_rows: list[dict] = []
    if slug == "L0372" and ar_text:
        schedule_rows = _extract_l0372_schedule_items(ar_text)

    # Merge by article_number (string key).
    merged: dict[str, dict] = {}

    for art in ar_articles:
        key = art["article_number"]
        merged[key] = {
            "article_number": key,
            "sort_order": art["sort_order"],
            "chapter": art.get("chapter_ar"),
            "section": None,
            "heading_ar": None,
            "heading_en": None,
            "text_ar": art.get("text_ar"),
            "text_en": None,
        }

    for art in en_articles:
        key = art["article_number"]
        if key in merged:
            merged[key]["text_en"] = art.get("text_en")
            if not merged[key]["chapter"]:
                merged[key]["chapter"] = art.get("chapter_en")
        else:
            merged[key] = {
                "article_number": key,
                "sort_order": art["sort_order"],
                "chapter": art.get("chapter_en"),
                "section": None,
                "heading_ar": None,
                "heading_en": None,
                "text_ar": None,
                "text_en": art.get("text_en"),
            }

    for art in schedule_rows:
        key = art["article_number"]
        merged[key] = {
            "article_number": key,
            "sort_order": art["sort_order"],
            "chapter": art.get("_chapter_ar"),
            "section": None,
            "heading_ar": art.get("_heading_ar"),
            "heading_en": None,
            "text_ar": art.get("text_ar"),
            "text_en": None,
        }

    if not merged:
        # UNPARSED fallback — include full text so nothing is silently dropped.
        notes.append("parser produced zero articles; UNPARSED row created")
        merged["UNPARSED"] = {
            "article_number": "UNPARSED",
            "sort_order": 0,
            "chapter": None,
            "section": None,
            "heading_ar": None,
            "heading_en": None,
            "text_ar": ar_text,
            "text_en": en_text,
        }

    # Stable sort by sort_order
    articles = sorted(merged.values(), key=lambda a: (a["sort_order"], a["article_number"]))

    notes.append(
        f"ar_articles={len(ar_articles)} en_articles={len(en_articles)} "
        f"schedule_rows={len(schedule_rows)} merged={len(articles)}"
    )
    return articles, notes


def main() -> None:
    # Collect per-slug text
    texts: dict[str, dict] = {}
    for lang_folder, lang_key in [("Arabic", "ar"), ("English", "en")]:
        folder = LAWS_ROOT / lang_folder
        if not folder.exists():
            print(f"skip — missing {folder}")
            continue
        for entry in sorted(folder.iterdir()):
            if not entry.is_file():
                continue
            slug = slug_from_filename(entry.name)
            if slug is None:
                continue
            row = texts.setdefault(slug, {"ar": None, "en": None, "files": []})
            try:
                if entry.suffix.lower() == ".docx":
                    body = extract_docx(entry)
                elif entry.suffix.lower() in {".html", ".htm"}:
                    body = extract_html(entry)
                else:
                    continue
            except Exception as exc:
                print(f"  ! {entry.name}: {exc}")
                continue
            row[lang_key] = body
            row["files"].append(entry.name)

    out: dict[str, dict] = {}
    for slug, bundle in sorted(texts.items()):
        articles, notes = parse_slug(slug, bundle.get("ar"), bundle.get("en"))
        out[slug] = {"files": bundle["files"], "notes": notes, "articles": articles}
        unparsed = any(a["article_number"] == "UNPARSED" for a in articles)
        flag = "  [UNPARSED]" if unparsed else ""
        print(f"  {slug:10s} articles={len(articles):4d}{flag}  notes={'; '.join(notes)}")

    # Strip private fields before dumping
    for slug, bundle in out.items():
        for a in bundle["articles"]:
            for k in list(a.keys()):
                if k.startswith("_"):
                    del a[k]

    OUT_FILE.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nWrote {OUT_FILE} — {len(out)} slugs")


if __name__ == "__main__":
    main()
