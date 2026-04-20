"""
extract_laws.py
---------------
One-time helper that walks C:\\Users\\sahwa\\Documents\\projects\\bahrain-laws\\
{Arabic,English} and extracts plain-text bodies of every instrument we have on
disk. Outputs a JSON file (scripts/_extracted_texts.json) keyed by slug with
{ "ar": "...", "en": "...", "ar_bytes": int, "en_bytes": int, "notes": [...] }.

Run from the sahwanlaw repo root:
    python scripts/extract_laws.py

Dependencies: python-docx, beautifulsoup4
"""
from __future__ import annotations

import json
import os
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document

LAWS_ROOT = Path(r"C:\Users\sahwa\Documents\projects\bahrain-laws")
OUT_FILE = Path(__file__).resolve().parent / "_extracted_texts.json"

# Filename prefix -> canonical slug (LLOC slug, lowercase here for matching).
# We keep the file's leading slug code (K0915, L0372, RJIW3020, etc.) as the
# canonical key.
SLUG_RE = re.compile(r"^([A-Z]+\d+)_")


def slug_from_filename(name: str) -> str | None:
    m = SLUG_RE.match(name)
    return m.group(1) if m else None


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # Also pull tables (LLOC consolidated docx occasionally uses tables for
    # amendment headers).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


# Selectors that match LLOC navigation chrome we want to drop.
LLOC_DROP_SELECTORS = [
    "header", "footer", "nav", "script", "style", "noscript",
    ".header", ".footer", ".navbar", ".menu", ".sidebar", ".breadcrumb",
    ".search", ".social", ".language", ".top-bar",
    "#header", "#footer", "#nav", "#menu", "#sidebar",
]


def extract_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")

    # Prefer the densest text container if it exists.
    main = (
        soup.find("div", id="legislationtextcontent")
        or soup.find("div", class_="legislationtextcontent")
        or soup.find("div", id="legislationtext")
        or soup.find("div", class_="legislationtext")
        or soup.find("main")
        or soup.find("article")
        or soup.body
        or soup
    )

    # Drop chrome.
    for sel in LLOC_DROP_SELECTORS:
        for el in main.select(sel):
            el.decompose()

    text = main.get_text("\n", strip=True)
    # Collapse blank-line runs.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def main() -> None:
    out: dict[str, dict] = {}

    for lang_folder, lang_key in [("Arabic", "ar"), ("English", "en")]:
        folder = LAWS_ROOT / lang_folder
        if not folder.exists():
            print(f"  skip — missing folder {folder}")
            continue
        for entry in sorted(folder.iterdir()):
            if not entry.is_file():
                continue
            slug = slug_from_filename(entry.name)
            if slug is None:
                print(f"  ?? cannot parse slug from {entry.name}")
                continue

            row = out.setdefault(slug, {"ar": None, "en": None, "ar_bytes": 0, "en_bytes": 0, "notes": []})
            try:
                if entry.suffix.lower() == ".docx":
                    body = extract_docx(entry)
                elif entry.suffix.lower() in {".html", ".htm"}:
                    body = extract_html(entry)
                else:
                    row["notes"].append(f"unsupported extension {entry.suffix} for {entry.name}")
                    continue
            except Exception as exc:  # noqa: BLE001
                row["notes"].append(f"extraction failure {entry.name}: {exc}")
                continue

            row[lang_key] = body
            row[f"{lang_key}_bytes"] = len(body.encode("utf-8"))
            row["notes"].append(f"{lang_key} <- {entry.name} ({len(body)} chars)")

    OUT_FILE.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {OUT_FILE} — {len(out)} slugs")
    for slug, row in sorted(out.items()):
        ar_len = len(row.get("ar") or "")
        en_len = len(row.get("en") or "")
        print(f"  {slug:8s} ar={ar_len:>7d} chars  en={en_len:>7d} chars")


if __name__ == "__main__":
    main()
